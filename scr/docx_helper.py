from docx.document import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.styles.style import BaseStyle
from docx.styles.styles import Styles
from settings import Settings

def export(data, document: Document, settings: Settings, finall: bool = False) -> None:
    print("Started exporintg")

    for paragraph in document.paragraphs:
        print(f"Trying to export: {paragraph.text.split(sep="\n")[0]}")

        for window in data:
            if paragraph.text.split(sep="\n")[0] == f"{settings("prefix", "zt")}-{window}":
                print(f"Exporting {window}")

                paragraph.text = f"{settings("prefix", "zt")}-{window}\n" if not finall else ""

                if settings("export_style", "internal") == "internal":
                    styles: Styles = document.styles
                    if "upper" not in styles:
                        upper: BaseStyle = styles.add_style("upper", WD_STYLE_TYPE.PARAGRAPH)
                        upper.quick_style = True
                        upper.base_style = styles['Normal'] #type:ignore
                        upper.font.italic = True #type:ignore
                        upper.font.size = Pt(11) #type:ignore

                    paragraph.style = "upper"
                
                elif settings("export_style", "internal") == "custom":
                    paragraph.style = str(settings("custom_sytel", "upper"))
                
                for index, line in enumerate(data[window]):
                    for letter in line:
                        match letter[1]:
                            case "n":
                                paragraph.add_run(letter[0])
                            case "d":
                                paragraph.add_run(letter[0]).font.subscript = True
                            case "u":
                                paragraph.add_run(letter[0]).font.superscript = True
                
                    if index != len(data[window]) - 1:
                        paragraph.add_run("\n")
    
    document.save(str(settings("docx", settings.chose_file(settings.file_types["docx"]))))
    print("Export done")

def docx_import(document: Document, settings) -> dict[str, list[list[tuple[str, str]]]]:
    print("Started importing...")
    data: dict[str, list[list[tuple[str, str]]]] = {}

    for paragraph in document.paragraphs:
        if paragraph.text.split("\n")[0].split("-")[0] == settings("prefix", "zt"):
            window: str = paragraph.text.split("\n")[0].split("-")[1]
            data[window] = []
            print(f"Importing window: {window}")
            
            start_index: int = 0
            runs: bool = True
            while runs:
                for i in paragraph.runs[start_index].text:
                    if i == "\n":
                        runs = False
                        break
                else:
                    start_index += 1
            
            had_line: bool = False
            for run in paragraph.runs[start_index:]:
                for letter in run.text:
                    if letter == "\n":
                        data[window].append([])
                        had_line = True
                        continue

                    if not had_line:
                        continue

                    if run.font.subscript:
                        data[window][0].append((letter, "d"))
                    
                    elif run.font.superscript:
                        data[window][0].append((letter, "u"))
                    
                    else:
                        data[window][0].append((letter, "n"))
    
    print("Import done")
    return data