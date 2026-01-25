from logging.config import valid_ident
from typing import Callable, Optional
import customtkinter as ctk
import keyboard
from string import ascii_lowercase
import json
from docx import Document

from settings import Settings, SettingsSetter
from window import Window
from window_helepr import WindowSwitcher, WindowLink
import docx_helper
from ex_bar import exFrame

type windows_t = dict[str, list[list[tuple[str, str]]]]
type window_t = list[list[tuple[str, str]]]

class App(ctk.CTk):
    def __init__(self) -> None:
        super().__init__()

        self.settings: Settings = Settings()

        self.geometry("700x330")
        self.title("Typer")
        self.config(bg="black")

        self.down_index_smart: bool = False
        self.down_index: bool = False
        self.up_index: bool = False
        self.up_index_smart: bool = False

        self.windows_bar_frame: WindowSwitcher = WindowSwitcher(self)
        # self.exbar: exBar = exBar(self)

        # To load window data
        self.load()

        self.create_hotkeys()

        self.windows_bar_frame.set_window(list(self.windows_bar_frame.windows)[0])("")
        self.mainloop()
    
    def create_hotkeys(self, from_ex = False) -> None:
        # To create hotkeys many times
        self.hotkeys: list[Callable] = []

        self.hotkeys.append(keyboard.add_hotkey("right", callback=self.on_direction_click_side(1)))
        self.hotkeys.append(keyboard.add_hotkey("left", callback=self.on_direction_click_side(-1)))
        self.hotkeys.append(keyboard.add_hotkey("up", callback=self.on_direction_click_height(-1)))
        self.hotkeys.append(keyboard.add_hotkey("down", callback=self.on_direction_click_height(1)))
        
        self.hotkeys.append(keyboard.add_hotkey("backspace", callback=self.delete_char))
        self.hotkeys.append(keyboard.add_hotkey("enter", callback=self.new_line))

        for i in ascii_lowercase.removeprefix("") + "=-/.":
            self.hotkeys.append(keyboard.add_hotkey(i, callback=self.type_key(i, "n")))
            self.hotkeys.append(keyboard.add_hotkey(f"ctrl+{i}", callback=self.type_key(i, "d")))
            self.hotkeys.append(keyboard.add_hotkey(f"alt+{i}", callback=self.type_key(i, "u")))
        
        for i in ascii_lowercase.upper():
            self.hotkeys.append(keyboard.add_hotkey(f"shift+{i}", callback=self.type_key(i, "n")))
            self.hotkeys.append(keyboard.add_hotkey(f"ctrl+shift+{i}", callback=self.type_key(i, "d")))
            self.hotkeys.append(keyboard.add_hotkey(f"alt+shift+{i}", callback=self.type_key(i, "u")))

        # Each letter hotkey needs to have ctrl - "d" and alt - "u" prefexies for indexing
        for i in ("n", "d", "u"):
            if i == "n":
                prefix: str = ""
            elif i == "d":
                prefix: str = "ctrl+"
            else:
                prefix: str = "alt+"

            self.hotkeys.append(keyboard.add_hotkey(prefix + "space", callback=self.type_key(" ", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "1", callback=self.type_key("+", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "altgr+-", callback=self.type_key("*", i)))

            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+é", callback=self.type_key("0", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+1", callback=self.type_key("1", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+ě", callback=self.type_key("2", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+š", callback=self.type_key("3", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+č", callback=self.type_key("4", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+ř", callback=self.type_key("5", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+ž", callback=self.type_key("6", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+ý", callback=self.type_key("7", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+á", callback=self.type_key("8", i)))
            self.hotkeys.append(keyboard.add_hotkey(prefix + "shift+í", callback=self.type_key("9", i)))
        
        # Definesm shortcuts for using
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+s", callback=self.save))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+c", callback=self.close_window))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+p", callback=self.change_settings))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+shift+s", callback=self.new_save))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+shift+n", callback=self.new_file))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+l", callback=self.load))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+shift+l", callback=self.force_load))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+n", callback=self.new_window))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+e", callback=self.export))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+f", callback=self.export_finall))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+i", callback=self.docx_import))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+right", callback=self.on_click_change_window(1)))
        self.hotkeys.append(keyboard.add_hotkey("ctrl+alt+left", callback=self.on_click_change_window(-1)))

        self.hotkeys.append(keyboard.add_hotkey("tab", callback=self.show))

        self.hotkeys.append(keyboard.add_hotkey("ctrl", callback=self.set_down_index))
        self.hotkeys.append(keyboard.add_hotkey("alt", callback=self.set_up_index))

    def clear_hotkeys(self) -> None:
        for i in self.hotkeys:
            keyboard.remove_hotkey(i)
    
    def show(self) -> None:
        self.clear_hotkeys()
        self.exbar = exFrame(self, self.un_show)
    
    def un_show(self) -> None:
        print(self.exbar.ret)
        self.exbar.place_forget()
        for letter in self.exbar.ret:
            self.type_key(letter[0], letter[1], ignore_keyboard = True)()
        self.after(60, self.windows_bar_frame.active_window().rerender_line, self.windows_bar_frame.active_window().line)
        # self.after(30, self.on_direction_click_height(-1))

        self.create_hotkeys()
    
    def export(self) -> None:
        # Exports to word document - leavs export marks
        self.save()
        docx_helper.export(self.save_data, Document(str(self.settings("docx", self.settings.chose_file(self.settings.file_types["docx"])))), self.settings)
    
    def export_finall(self) -> None:
        # Exports to word document - delets export marks
        self.save()
        docx_helper.export(self.save_data, Document(str(self.settings("docx", self.settings.chose_file(self.settings.file_types["docx"])))), self.settings, finall=True)
    
    def docx_import(self) -> None:
        # Imports from word document
        self.save()
        data: windows_t = docx_helper.docx_import(Document(str(self.settings("docx", self.settings.chose_file(self.settings.file_types["docx"])))), self.settings)
        for window in data:
            self.create_window(window, data[window])

    def on_click_change_window(self, val: int) -> Callable[[], None]:
        # Switchis windows with shortcut ctrl+alt+ right | left
        def run() -> None:

            # Translets dictionary into list for index aproching
            ac: str = self.windows_bar_frame.a_window
            len_windows: int = len(self.windows_bar_frame.windows)
            a_window: int = list(self.windows_bar_frame.windows).index(ac)

            if a_window + val >= len_windows:
                a_window = 0
            elif a_window + val == -1:
                a_window = len_windows - 1
            else:
                a_window += val
            
            self.windows_bar_frame.set_window(list(self.windows_bar_frame.windows)[a_window])(None)

        return run
    
    def new_window(self) -> None:
        # Creates new windw
        # Clears hotkeys to not enter new line at the end
        self.clear_hotkeys()

        title: str = str(ctk.CTkInputDialog(text="New window name:").get_input())
        self.create_window(title=title, content=None)
        
        self.after(10, lambda: self.create_hotkeys())
    
    def new_line(self) -> None:
        self.windows_bar_frame.active_window().new_line()
    
    def delete_char(self) -> None:
        self.windows_bar_frame.active_window().delete_char()
    
    def on_direction_click_side(self, val: int) -> Callable[[], None]:
        # Forwarding of function for moving cursor horizontaly
        def run() -> None:
            self.windows_bar_frame.active_window().on_direction_click_side(val)()
        
        return run

    def on_direction_click_height(self, val: int) -> Callable[[], None]:
        # Forwarding of function for moving cursor verticaly
        def run() -> None:
            self.windows_bar_frame.active_window().on_direction_click_height(val)()
        
        return run
    
    def type_key(self, key: str, type2: str, ignore_keyboard: bool = False) -> Callable[[], None]:
        def run() -> None:
            type: str = type2

            # Togle indexing
            if self.settings("index_mode", "toggle") == "hold":
                if self.up_index:
                    type = "u"
                
                if self.down_index:
                    type = "d"
            
            # Smart indexing
            if self.settings("smart_index", "off") == "on":
                if self.up_index_smart:
                    type = "u"
                
                if self.down_index_smart:
                    type = "d"
            
            # Activation of smart indexing
            if key in ascii_lowercase or key in ascii_lowercase.upper():
                self.down_index_smart = True
            
            # End of smart indexing
            elif key == " ":
                self.down_index_smart = False
                type = "n"
            
            self.windows_bar_frame.active_window().type_key(key, type, ignor_keyboard=ignore_keyboard)()
            print(key, type)
        
        return run

    def create_window(self, title: str, content: Optional[window_t]) -> None:
        self.windows_bar_frame.add_window(WindowLink(self.windows_bar_frame, Window(self, content=content), title=title if title != "" else "Not defined"))
    
    def save(self) -> None:
        # Saves and it is used to generate self.save_data multiple times
        print("startings to save...")
        self.save_data: windows_t = {}

        for i in self.windows_bar_frame.windows:
            print(f"saving: {i}")
            self.save_data[self.windows_bar_frame.windows[i].title] = self.windows_bar_frame.windows[i].window.save()
            with open(f"{self.windows_bar_frame.windows[i].title}.json", "w") as file:
                json.dump(self.save_data[self.windows_bar_frame.windows[i].title], file)
        
        with open(str(self.settings("save", self.settings.chose_file(self.settings.file_types["json"]))), "w") as file:
            json.dump(self.save_data, file, indent=4)
        
        print("saving done")
    
    def new_save(self) -> None:
        # Creates new save file with the same content as the old one
        self.save()

        self.settings.data["save"] = self.settings.create_save().removesuffix(".json") + ".json"

        with open(self.settings.data["save"], "w") as file:
            json.dump(self.save_data, file, indent=4)
        print("Created new savefile")
    
    def new_file(self) -> None:
        # Creates brand new save file with only one empty window
        self.save()
        self.reload()

        title: str = str(ctk.CTkInputDialog(text="New window name:").get_input())
        self.settings.data["save"] = self.settings.create_save().removesuffix(".json") + ".json"
        with open(self.settings.data["save"], "w") as file:
            json.dump({title: [[]]}, file, indent=4)
        
        self.load()
    
    def load(self) -> None:
        print("started loading...")
        with open(str(self.settings("save", self.settings.chose_file(self.settings.file_types["json"]))), "r") as file:
            self.save_data: windows_t = json.load(file)
        
        for i in self.save_data:
            print(f"loading: {i}")
            self.create_window(i, self.save_data[i])
        
        print("loading done")
    
    def force_load(self) -> None:
        # Closes all windows befor loading
        print("started loading...")
        self.reload()

        with open(str(self.settings("save", self.settings.chose_file(self.settings.file_types["json"]), True)), "r") as file:
            self.save_data: windows_t = json.load(file)
        
        for i in self.save_data:
            print(f"loading: {i}")
            self.create_window(i, self.save_data[i])
        
        print("loading done")
    
    def reload(self) -> None:
        # Closes all windows
        print("Saterted reloading")
        for i in self.windows_bar_frame.windows.copy():
            self.windows_bar_frame.remove_window(i)
        print("Reloading done")
    
    def change_settings(self) -> None:
        # Opens settings menu
        self.setter: SettingsSetter = SettingsSetter(self.settings, self)
    
    def close_window(self) -> None:
        self.windows_bar_frame.remove_window(self.windows_bar_frame.a_window)
    
    def set_up_index(self) -> None:
        # Togle indexis + turning off smart indexing
        self.up_index = not self.up_index
        self.down_index_smart = False
        self.up_index_smart = False
    
    def set_down_index(self) -> None:
        # Togle indexis + turning off smart indexing
        self.down_index = not self.down_index
        self.down_index_smart = False
        self.up_index_smart = False


if __name__ == "__main__":
    App()